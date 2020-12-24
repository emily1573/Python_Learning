from docx import Document # 创建文档
from docx.oxml.ns import qn # 中文
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # 段落
from docx.shared import Pt,RGBColor,Mm,Cm # 大小 磅数 字号

from openpyxl import load_workbook
# 我们要设置三个模板
# 模板1
def template1(name):
    # 创建嘉宾模板 男士
    word_document = Document() # 创建word文档对象
    word_document.styles['Normal'].font.name = u'微软雅黑' # 正文/标题1/标题2 （英文)
    word_document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑') # 中文

    # -------段落创建P1--------
    p1 = word_document.add_paragraph() # 向word添加段落
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # 段落剧中对齐
    # -------添加文字----------
    run_text_1 = p1.add_run('万门大学客户大会')
    run_text_1.font.size = Pt(20)

    # --------段落创建P2--------
    p2 = word_document.add_paragraph() # 向word添加段落
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # 段落居中对齐
    # --------添加文字-----------
    run_text_2 = p2.add_run('邀请函')
    run_text_2.font.size = Pt(36)

    # ---------段落创建p customer-------
    p_customer = word_document.add_paragraph() # 向word添加段落
    p_customer.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT # 段落左对齐
    # ----------添加文字----------------
    run_text_c = p_customer.add_run('尊敬的 %s 先生' % name)
    run_text_c.font.size = Pt(20)

    # ----------创建段落p3--------------
    p3 = word_document.add_paragraph() # 向word添加段落
    # -----------添加文字---------------
    run_text_3 = p3.add_run('    诚邀您参加由万门大学组织的客户答谢会，请于2020年10月1日在万门大学总部准时出席，谢谢！')
    run_text_3.font.size = Pt(16)

    # ----------段落创建P4 公司logo图片----------
    p4 = word_document.add_paragraph() # 向word添加段落
    p4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # 段落居中对齐

    # ----------添加文字--------------------
    run_text_4 = p4.add_run()
    run_text_4.add_picture('wmlogo.jpeg' , width=Mm(80))

    word_document.save('VIP/%s.docx' % name)
# 模板2
def template2(name):
    # 创建嘉宾模板 女士
    word_document = Document()  # 创建word文档对象
    word_document.styles['Normal'].font.name = u'微软雅黑'  # 正文/标题1/标题2 （英文)
    word_document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 中文

    # -------段落创建P1--------
    p1 = word_document.add_paragraph()  # 向word添加段落
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 段落剧中对齐
    # -------添加文字----------
    run_text_1 = p1.add_run('万门大学客户大会')
    run_text_1.font.size = Pt(20)

    # --------段落创建P2--------
    p2 = word_document.add_paragraph()  # 向word添加段落
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 段落居中对齐
    # --------添加文字-----------
    run_text_2 = p2.add_run('邀请函')
    run_text_2.font.size = Pt(36)

    # ---------段落创建p customer-------
    p_customer = word_document.add_paragraph()  # 向word添加段落
    p_customer.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 段落左对齐
    # ----------添加文字----------------
    run_text_c = p_customer.add_run('尊敬的 %s 女士' % name)
    run_text_c.font.size = Pt(20)

    # ----------创建段落p3--------------
    p3 = word_document.add_paragraph()  # 向word添加段落
    # -----------添加文字---------------
    run_text_3 = p3.add_run('    诚邀您参加由万门大学组织的客户答谢会，请于2020年10月1日在万门大学总部准时出席，谢谢！')
    run_text_3.font.size = Pt(16)

    # ----------段落创建P4 公司logo图片----------
    p4 = word_document.add_paragraph()  # 向word添加段落
    p4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 段落居中对齐

    # ----------添加文字--------------------
    run_text_4 = p4.add_run()
    run_text_4.add_picture('wmlogo.jpeg', width=Mm(80))

    word_document.save('VIP/%s.docx' % name)
# 模板3
def template3(name):
    # 创建媒体模板
    word_document = Document()  # 创建word文档对象
    word_document.styles['Normal'].font.name = u'微软雅黑'  # 正文/标题1/标题2 （英文)
    word_document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 中文

    # -------段落创建P1--------
    p1 = word_document.add_paragraph()  # 向word添加段落
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 段落剧中对齐
    # -------添加文字----------
    run_text_1 = p1.add_run('万门大学客户大会')
    run_text_1.font.name = u'宋体'
    run_text_1._element.rPr.rFonts.set(qn('w:eastAsia'), u's宋体')
    run_text_1.font.size = Pt(20)

    # --------段落创建P2--------
    p2 = word_document.add_paragraph()  # 向word添加段落
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 段落居中对齐
    # --------添加文字-----------
    run_text_2 = p2.add_run('邀请函')
    run_text_2.font.size = Pt(36)

    # ---------段落创建p customer-------
    p_customer = word_document.add_paragraph()  # 向word添加段落
    p_customer.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 段落左对齐
    # ----------添加文字----------------
    run_text_c = p_customer.add_run('尊敬的 %s 老师' % name)
    run_text_c.font.size = Pt(20)
    run_text_c.font.color.rgb = RGBColor(177, 23, 79) # 注意这里是改变颜色的

    # ----------创建段落p3--------------
    p3 = word_document.add_paragraph()  # 向word添加段落
    # -----------添加文字---------------
    run_text_3 = p3.add_run('    诚邀您参加由万门大学组织的客户答谢会，请于2020年10月1日在万门大学总部准时出席，谢谢！')
    run_text_3.font.size = Pt(16)

    # ----------段落创建P4 公司logo图片----------
    p4 = word_document.add_paragraph()  # 向word添加段落
    p4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 段落居中对齐

    # ----------添加文字--------------------
    run_text_4 = p4.add_run()
    run_text_4.add_picture('wmlogo.jpeg', width=Mm(80))

    word_document.save('Media/%s.docx' % name)

# ******以上为模板 以下为控制 读出数 然后 分模板进行创建word
# 主控
def main_creater():
    wb = load_workbook('VIPLIST.xlsx')
    ws = wb.worksheets[0]
    for r in list(ws.iter_rows())[1:]:
        row = [ data.value for data in r ] # 注意这个推到列表非常重要
        if row[1] == '男' and row[2] == '嘉宾':
            # 男士称呼 保存嘉宾目录
            template1(row[0])

            pass
        elif row[1] == '女' and row[2] == '嘉宾':
            # 女士称呼 保存嘉宾目录
            template2(row[0])
            pass
        elif row[2] == '媒体':
            # 媒体称呼 保存媒体目录
            template3(row[0])
            pass

if __name__ == '__main__':
    main_creater()



